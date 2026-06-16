"""
Converts PROJECT_REPORT.md into a formatted Word document.
Run:  python "project docs/build_report.py"
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── paths ─────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
MD_FILE    = SCRIPT_DIR / "PROJECT_REPORT.md"
OUT_FILE   = SCRIPT_DIR / "AI_Fake_News_Detector_Project_Report.docx"

# ── helpers ───────────────────────────────────────────────────────────────────

def set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


def para_space(doc, before=0, after=6):
    """Tiny spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)


# ── styles ────────────────────────────────────────────────────────────────────

def configure_styles(doc):
    styles = doc.styles

    # Normal / body
    n = styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.paragraph_format.space_after  = Pt(6)
    n.paragraph_format.space_before = Pt(0)
    n.paragraph_format.line_spacing = Pt(14)

    # Heading 1  — chapter title
    h1 = styles["Heading 1"]
    h1.font.name    = "Calibri"
    h1.font.size    = Pt(16)
    h1.font.bold    = True
    h1.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2  — section title
    h2 = styles["Heading 2"]
    h2.font.name    = "Calibri"
    h2.font.size    = Pt(13)
    h2.font.bold    = True
    h2.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3  — sub-section
    h3 = styles["Heading 3"]
    h3.font.name  = "Calibri"
    h3.font.size  = Pt(11)
    h3.font.bold  = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0x2E, 0x54, 0x96)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(3)
    h3.paragraph_format.keep_with_next = True

    # Code / monospace
    if "Code" not in [s.name for s in styles]:
        code_style = styles.add_style("Code", 1)
    else:
        code_style = styles["Code"]
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after  = Pt(4)
    code_style.paragraph_format.left_indent  = Inches(0.3)

    return styles


# ── title page ────────────────────────────────────────────────────────────────

def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("RV College of Engineering")
    run.font.name  = "Calibri"
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Department of Master of Computer Applications (MCA)")
    r2.font.name = "Calibri"
    r2.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("AI-Based Fake News Detection Using\nLarge Language Models")
    r3.font.name  = "Calibri"
    r3.font.size  = Pt(20)
    r3.font.bold  = True
    r3.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Project Report — 2025-26")
    r4.font.name  = "Calibri"
    r4.font.size  = Pt(12)
    r4.font.italic = True

    for _ in range(6):
        doc.add_paragraph()

    for label, value in [
        ("Student", "Sivabalan"),
        ("Institution", "RV College of Engineering, Bengaluru"),
        ("Department", "MCA"),
        ("Academic Year", "2025-26"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}: ")
        r.font.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)

    doc.add_page_break()


# ── inline formatting ─────────────────────────────────────────────────────────

BOLD_RE   = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"\*(.+?)\*(?!\*)")
CODE_RE   = re.compile(r"`(.+?)`")


def add_inline(para, text):
    """Add text with **bold**, *italic*, and `code` spans."""
    pos = 0
    combined = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*(?!\*)|`(.+?)`")
    for m in combined.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        raw = m.group(0)
        if raw.startswith("**"):
            r = para.add_run(m.group(1))
            r.bold = True
        elif raw.startswith("`"):
            r = para.add_run(m.group(3))
            r.font.name = "Courier New"
            r.font.size = Pt(9.5)
        else:
            r = para.add_run(m.group(2))
            r.italic = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


# ── table renderer ────────────────────────────────────────────────────────────

def render_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # distribute column widths evenly
    total_width = Inches(6.0)
    col_width   = total_width / col_count

    for row_idx, row_data in enumerate(rows):
        row_obj = table.add_row()
        row_obj.height = Cm(0.7)
        for ci, cell_text in enumerate(row_data):
            cell = row_obj.cells[ci]
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)

            if row_idx == 0:
                set_cell_bg(cell, "1F3564")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(cell_text.strip())
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
                r.font.name = "Calibri"
            else:
                bg = "EBF0FA" if row_idx % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(p, cell_text.strip())
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── code block renderer ───────────────────────────────────────────────────────

def render_code_block(doc, lines):
    # shaded box for code
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.0)
    set_cell_bg(cell, "F4F4F4")
    set_cell_border(cell)

    # clear existing paragraph
    for para in list(cell.paragraphs):
        p_elem = para._element
        p_elem.getparent().remove(p_elem)

    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── main parser ───────────────────────────────────────────────────────────────

def parse_and_build(doc, md_text):
    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── horizontal rule ──
        if re.match(r"^-{3,}$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── heading 1 (# ) ──
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # ── heading 2 (## ) ──
        if line.startswith("## ") and not line.startswith("### "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # ── heading 3 (### ) ──
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # ── fenced code block ``` ──
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            render_code_block(doc, code_lines)
            continue

        # ── markdown table ──
        if "|" in line and line.strip().startswith("|"):
            table_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                raw = lines[i].strip().strip("|")
                cells = [c.strip() for c in raw.split("|")]
                # skip separator row (---|---|---)
                if not all(re.match(r"^[-: ]+$", c) for c in cells if c):
                    table_rows.append(cells)
                i += 1
            render_table(doc, table_rows)
            continue

        # ── bullet list ──
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            add_inline(p, line[2:].strip())
            i += 1
            continue

        # ── numbered list ──
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            content = re.sub(r"^\d+\. ", "", line)
            add_inline(p, content.strip())
            i += 1
            continue

        # ── blank line ──
        if not line.strip():
            i += 1
            continue

        # ── normal paragraph ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, line.strip())
        i += 1


# ── entry point ───────────────────────────────────────────────────────────────

def main():
    print(f"Reading  : {MD_FILE}")
    md_text = MD_FILE.read_text(encoding="utf-8")

    doc = Document()
    set_page_margins(doc)
    configure_styles(doc)
    add_title_page(doc)

    # running header/footer
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        r = hp.add_run("AI-Based Fake News Detection Using Large Language Models")
        r.font.name  = "Calibri"
        r.font.size  = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = fp.add_run("Department of MCA, RVCE    ")
        r2.font.name = "Calibri"
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        # page number field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        r3 = fp.add_run()
        r3._r.append(fldChar1)
        r3._r.append(instrText)
        r3._r.append(fldChar2)
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    parse_and_build(doc, md_text)

    doc.save(str(OUT_FILE))
    print(f"Saved    : {OUT_FILE}")
    print("Done.")


if __name__ == "__main__":
    main()
